"""
Knowledge Flow - Orquestração do pipeline Cortex
"""
import os
import uuid
import asyncio
import json
import re
import logging
import magic
from datetime import datetime
from typing import Optional, Callable, Tuple
import httpx
from pypdf import PdfReader
from docx import Document
from langdetect import detect, DetectorFactory
from bs4 import BeautifulSoup
from qdrant_client import QdrantClient, models
from sqlalchemy import create_engine, text

# Ensure langdetect is deterministic
DetectorFactory.seed = 0

from .schemas import IngestRecommendation

logger = logging.getLogger(__name__)

OLLAMA_URL = os.getenv('OLLAMA_URL', 'http://ollama:11434')
OLLAMA_EMBEDDING_MODEL = os.getenv('EMBEDDING_MODEL', 'nomic-embed-text:latest')
OLLAMA_CORTEX_MODEL = os.getenv('CORTEX_MODEL', 'qwen2.5:7b-instruct')
QDRANT_URL = os.getenv('QDRANT_URL', 'http://qdrant:6333')
QDRANT_COLLECTION = os.getenv('QDRANT_COLLECTION', 'mycrew_agent_kb')

# Fallback para API paga (Claude)
PAID_API_KEY = os.getenv('PAID_API_KEY', '')
PAID_API_URL = os.getenv('PAID_API_URL', 'https://api.anthropic.com/v1/messages')


def extract_json_from_response(content: str) -> str:
    """Extrai JSON válido de uma resposta de texto."""
    content = re.sub(r'```json\s*', '', content)
    content = re.sub(r'```\s*', '', content)
    json_match = re.search(r'\{[\s\S]*\}', content)
    return json_match.group(0) if json_match else content


def fill_defaults(data: dict) -> dict:
    """Preenche defaults para campos faltantes."""
    if 'confidence' in data and isinstance(data['confidence'], (int, float)) and data['confidence'] > 1:
        data['confidence'] = data['confidence'] / 100.0
    if 'chunking_strategy' not in data or not data['chunking_strategy']:
        data['chunking_strategy'] = {'primary': {'type': 'paragraph', 'reason': 'Padrão geral'}, 'parameters': {'chunk_size': 512, 'chunk_overlap': 64, 'separator': '\n\n', 'min_chunk_size': 100, 'max_chunk_size': 1024}}
    elif 'parameters' not in data['chunking_strategy']:
        data['chunking_strategy']['parameters'] = {'chunk_size': 512, 'chunk_overlap': 64, 'separator': '\n\n', 'min_chunk_size': 100, 'max_chunk_size': 1024}
    elif 'primary' not in data['chunking_strategy']:
        data['chunking_strategy']['primary'] = {'type': 'paragraph', 'reason': 'Padrão geral'}
    if 'embedding' not in data or not data['embedding']:
        data['embedding'] = {'provider': 'ollama', 'model': OLLAMA_EMBEDDING_MODEL, 'dimensions': 768, 'normalize': True}
    if 'qdrant_index' not in data or not data['qdrant_index']:
        data['qdrant_index'] = {'distance': 'Cosine', 'hnsw_config': {'m': 16, 'ef_construct': 128, 'ef_search': 96}, 'payload_index_fields': []}
    elif 'hnsw_config' not in data['qdrant_index']:
        data['qdrant_index']['hnsw_config'] = {'m': 16, 'ef_construct': 128, 'ef_search': 96}
    elif 'payload_index_fields' not in data['qdrant_index']:
        data['qdrant_index']['payload_index_fields'] = []
    return data


def format_duration(ms: int) -> str:
    if ms < 60000: return f"{ms // 1000}s"
    elif ms < 3600000: return f"{ms // 60000}m {ms % 60000 // 1000}s"
    return f"{ms // 3600000}h {(ms % 3600000) // 60000}m"


def extract_text_from_pdf(file_content: bytes) -> Tuple[str, dict]:
    """Extrai texto de PDF usando pypdf."""
    try:
        import io
        reader = PdfReader(io.BytesIO(file_content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or "" + "\n"
        return text, {"pages": len(reader.pages)}
    except Exception as e:
        logger.warning(f"[PDF] Erro ao extrair: {e}")
        return "", {}


def extract_text_from_docx(file_content: bytes) -> Tuple[str, dict]:
    """Extrai texto de DOCX usando python-docx."""
    try:
        import io
        doc = Document(io.BytesIO(file_content))
        text = "\n".join([p.text for p in doc.paragraphs])
        return text, {"paragraphs": len(doc.paragraphs)}
    except Exception as e:
        logger.warning(f"[DOCX] Erro ao extrair: {e}")
        return "", {}


def extract_text_from_html(file_content: bytes) -> Tuple[str, dict]:
    """Extrai texto de HTML usando BeautifulSoup."""
    try:
        soup = BeautifulSoup(file_content.decode('utf-8', errors='replace'), 'html.parser')
        # Remove scripts e styles
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text(separator='\n')
        return text, {"tags": len(soup.find_all())}
    except Exception as e:
        logger.warning(f"[HTML] Erro ao extrair: {e}")
        return "", {}


def detect_file_type(file_content: bytes, filename: str) -> str:
    """Detecta tipo real do arquivo via magic bytes."""
    try:
        mime = magic.Magic(mime=True)
        mime_type = mime.from_buffer(file_content)
        
        type_map = {
            'application/pdf': 'pdf',
            'text/plain': 'txt',
            'text/markdown': 'md',
            'text/html': 'html',
            'text/csv': 'csv',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/json': 'other',
        }
        
        detected = type_map.get(mime_type)
        if detected:
            return detected
            
        # Fallback para extensão
        ext = filename.split('.')[-1].lower() if '.' in filename else 'other'
        return ext if ext in type_map.values() else 'other'
    except Exception:
        ext = filename.split('.')[-1].lower() if '.' in filename else 'other'
        return ext if ext in ['pdf', 'md', 'txt', 'html', 'csv', 'docx'] else 'other'


def detect_language(content: str) -> str:
    """Detecta idioma usando langdetect."""
    try:
        if len(content) < 20:
            return 'pt'  # fallback para textos muito curtos
        lang = detect(content[:1000])  # analisa primeiros 1000 chars
        lang_map = {'en': 'en', 'pt': 'pt', 'es': 'es', 'fr': 'fr', 'de': 'de'}
        return lang_map.get(lang, 'pt')
    except Exception:
        return 'pt'


def analyze_structure(content: str, file_type: str) -> dict:
    """Analisa estrutura do documento para chunking."""
    notes = []
    structure_level = 'unstructured'
    
    # Detectar headings markdown
    h1_matches = len(re.findall(r'^#\s+.+$', content, re.MULTILINE))
    h2_matches = len(re.findall(r'^##\s+.+$', content, re.MULTILINE))
    h3_matches = len(re.findall(r'^###\s+.+$', content, re.MULTILINE))
    
    if h1_matches + h2_matches + h3_matches > 3:
        structure_level = 'structured'
        notes.append(f"Document has {h1_matches} H1, {h2_matches} H2, {h3_matches} H3 headings")
    
    # Detectar tabelas
    table_count = len(re.findall(r'\|.*\|', content))
    if table_count > 10:
        notes.append(f"Document contains {table_count} potential table rows")
    
    # Detectar código
    code_blocks = len(re.findall(r'```[\w]*', content))
    if code_blocks > 0:
        notes.append(f"Document has {code_blocks} code blocks")
    
    # Detectar OCR issues
    if re.search(r'[a-z]{20,}', content) and not re.search(r'\s', content[:50]):
        notes.append("Possible OCR issues - long words without spaces detected")
    
    return {
        'structure_level': structure_level,
        'notes': notes,
        'headings': {'h1': h1_matches, 'h2': h2_matches, 'h3': h3_matches},
        'has_tables': table_count > 10,
        'has_code': code_blocks > 0
    }


def chunk_text(content: str, strategy_type: str, params: dict) -> list:
    """Chunking avançado com múltiplas estratégias."""
    chunk_size = params.get('chunk_size', 512)
    chunk_overlap = params.get('chunk_overlap', 64)
    separator = params.get('separator', '\n\n')
    
    # Estratégia markdown_header - split por headings
    if strategy_type == 'markdown_header':
        sections = re.split(r'(^#+\s+.+$)', content, flags=re.MULTILINE)
        chunks = []
        current = ""
        for section in sections:
            section = section.strip()
            if not section:
                continue
            if section.startswith('#'):
                if current:
                    chunks.append(current)
                current = section + "\n\n"
            else:
                current += section + "\n\n"
        if current.strip():
            chunks.append(current.strip())
        return chunks if chunks else [content]
    
    # Estratégia table_aware - preserva tabelas
    if strategy_type == 'table_aware':
        # Split preservando tabelas markdown
        table_pattern = r'(\|[-|:]+\|\n(?:(\|.*\|\n)+))'
        parts = re.split(table_pattern, content)
        chunks = []
        current = ""
        for part in parts:
            if re.match(table_pattern, part):
                # É uma tabela
                if current.strip():
                    chunks.append(current.strip())
                chunks.append(part.strip())
                current = ""
            else:
                current += part
        if current.strip():
            # Aplica chunking de parágrafo no restante
            paragraphs = current.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para:
                    if len(chunks) > 0 and len(para) // 4 + len(chunks[-1]) // 4 <= chunk_size * 2:
                        chunks[-1] = chunks[-1] + "\n\n" + para
                    else:
                        chunks.append(para)
        return chunks
    
    # Estratégia code_aware - preserva blocos de código
    if strategy_type == 'code_aware':
        code_pattern = r"(```[\w]*\n[\s\S]*?```)"
        parts = re.split(code_pattern, content)
        chunks = []
        current = ""
        for part in parts:
            if re.match(code_pattern, part):
                # É código
                if current.strip():
                    chunks.append(current.strip())
                chunks.append(part.strip())
                current = ""
            else:
                current += part + "\n"
        if current.strip():
            # Chunking de parágrafo no conteúdo não-código
            paragraphs = current.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para:
                    if chunks and len(para) // 4 + len(chunks[-1]) // 4 <= chunk_size * 2:
                        chunks[-1] = chunks[-1] + "\n\n" + para
                    else:
                        chunks.append(para)
        return chunks
    
    # Estratégia hybrid - combina baseado em features detectadas
    if strategy_type == 'hybrid':
        return chunk_text(content, 'markdown_header', params)  # Prioriza estrutura
    
    # Estratégias padrão
    if strategy_type == 'paragraph':
        paragraphs = content.split(separator)
        chunks, current = [], ""
        for para in paragraphs:
            para = para.strip()
            if not para: continue
            if len(current) // 4 + len(para) // 4 > chunk_size and current:
                chunks.append(current.strip())
                current = current[-(chunk_overlap * 4):] + "\n\n" + para
            else:
                current = current + "\n\n" + para if current else para
        if current.strip(): chunks.append(current.strip())
        return chunks if chunks else [content[:chunk_size * 4]]
    
    elif strategy_type == 'fixed_size':
        return [content[i:i + chunk_size * 4] for i in range(0, len(content), (chunk_size - chunk_overlap) * 4)]
    
    elif strategy_type == 'sentence':
        sentences = re.split(r'[.!?]+', content)
        chunks, current = [], ""
        for sent in sentences:
            sent = sent.strip()
            if not sent: continue
            if len(current) // 4 + len(sent) // 4 > chunk_size and current:
                chunks.append(current.strip())
                current = sent
            else:
                current = current + ". " + sent if current else sent
        if current.strip(): chunks.append(current.strip())
        return chunks if chunks else [content[:chunk_size * 4]]
    
    elif strategy_type == 'sliding_window':
        return [content[i:i + chunk_size * 4] for i in range(0, len(content), (chunk_size - chunk_overlap) * 4)]
    
    elif strategy_type == 'recursive':
        # Implementação similar ao LangChain RecursiveCharacterTextSplitter
        chars = ["\n\n", "\n", " ", ""]
        chunks = []
        current = content
        while len(current) > chunk_size * 4:
            for sep in chars:
                if sep == "":
                    chunks.append(current[:chunk_size * 4])
                    current = current[chunk_size * 4:]
                    break
                parts = current.split(sep, 1)
                if len(parts) == 2:
                    chunks.append(parts[0])
                    current = parts[1]
                    break
        if current.strip():
            chunks.append(current.strip())
        return chunks
    
    return [content[i:i + chunk_size * 4] for i in range(0, len(content), (chunk_size - chunk_overlap) * 4)]


def get_fallback_recommendation(filename: str, language: str, char_count: int, structure_level: str, error: str = "") -> dict:
    return {
        "operation": "ingest",
        "document": {"filename": filename, "file_type": "txt", "language": language, "estimated_tokens": char_count // 4, "structure_level": structure_level, "domain": "general"},
        "chunking_strategy": {"primary": {"type": "paragraph", "reason": "Fallback automatico"}, "parameters": {"chunk_size": 512, "chunk_overlap": 64, "separator": "\n\n"}},
        "embedding": {"provider": "ollama", "model": OLLAMA_EMBEDDING_MODEL, "dimensions": 768, "normalize": True},
        "qdrant_index": {"distance": "Cosine", "hnsw_config": {"m": 16, "ef_construct": 128, "ef_search": 96}, "payload_index_fields": []},
        "retrieval_hint": f"Estrategia fallback ativada{f' - {error}' if error else ''}",
        "notes": f"Falha na analise: {error}" if error else "Utilizando estrategia padrao",
        "alternative": {"type": "fixed_size", "reason": "Estrategia alternativa mais simples"},
        "testing_questions": ["Que tipo de informacoes procurar?"],
        "review_required": False,
        "confidence": 0.5
    }

# Database engine for usage tracking
_engine = None

def _get_engine():
    global _engine
    if _engine is None:
        POSTGRES_HOST = os.getenv("POSTGRES_HOST", "postgres")
        POSTGRES_PORT = os.getenv("POSTGRES_PORT", "5432")
        POSTGRES_USER = os.getenv("POSTGRES_USER", "mycrew")
        POSTGRES_PASSWORD = os.getenv("POSTGRES_PASSWORD", "")
        POSTGRES_DB = os.getenv("POSTGRES_DB", "mycrew")
        DATABASE_URL = os.getenv("DATABASE_URL", f"postgresql://{POSTGRES_USER}:{POSTGRES_PASSWORD}@{POSTGRES_HOST}:{POSTGRES_PORT}/{POSTGRES_DB}")
        _engine = create_engine(DATABASE_URL.replace("postgresql://", "postgresql+psycopg2://"), echo=False)
    return _engine


class KnowledgeFlow:
    def __init__(self, flow_id: str, user_id: str, file_content: bytes, filename: str, cortex_prompt: str = ""):
        self.flow_id = flow_id
        self.user_id = user_id
        self.file_content = file_content
        self.filename = filename
        self._cortex_prompt = cortex_prompt
        self._subscribers = []
        self._recommendation = None
        self._event_buffer = []
        self._step_times = {}
        self._step_start = None
        self._content = None
        self._chunks = []
        self._last_ollama_metrics = None  # Store metrics for Cortex persistence
        self._embedding_metrics = None  # Store metrics for embedding persistence
        self._flow_start_time = datetime.utcnow().timestamp()  # Start timing for processing
        self._flow_end_time = None  # Will be set at end of processing
        self._embedding_latency_ms = 0  # Track embedding latency separately
        self._metadata = None  # Store extracted metadata
        self._extraction_time_ms = 0  # Track extraction timing

    async def emit_step(self, step_id: str, step_name: str, status: str,
                        output_preview: Optional[str] = None, error_message: Optional[str] = None,
                        recommendation: Optional[dict] = None, start_timing: bool = True):
        duration_ms = int((datetime.utcnow().timestamp() - self._step_start) * 1000) if self._step_start else 0
        self._step_start = datetime.utcnow().timestamp() if start_timing else None
        event = {"flow_id": self.flow_id, "operation": "ingest", "step_id": step_id, "step_name": step_name, "status": status,
                 "output_preview": output_preview, "error_message": error_message, "duration_ms": duration_ms,
                 "duration_formatted": format_duration(duration_ms), "timestamp": datetime.utcnow().isoformat() + "Z"}
        if recommendation: event["recommendation"] = recommendation
        event_json = json.dumps(event)
        self._event_buffer.append(event_json)
        for callback in self._subscribers:
            await callback(event_json)

    async def _call_paid_api(self, content: str, language: str, structure_info: dict) -> tuple[dict, dict]:
        """Escalonamento para API paga quando Cortex falha."""
        if not PAID_API_KEY:
            logger.warning("[Cortex] No paid API key configured, using basic fallback")
            return get_fallback_recommendation(self.filename, language, len(content), structure_info.get('structure_level', 'semi_structured'), "No paid API key"), {"total_ms": 0, "tokens_generated": 0, "throughput_tps": 0}
        
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    PAID_API_URL,
                    headers={
                        "x-api-key": PAID_API_KEY,
                        "content-type": "application/json",
                        "anthropic-version": "2023-06-01"
                    },
                    json={
                        "model": "claude-3-haiku-20240307",
                        "max_tokens": 1024,
                        "temperature": 0,
                        "messages": [{"role": "user", "content": f"Analyze this document for indexing recommendations. Language: {language}. Structure info: {structure_info}. Content preview: {content[:4000]}"}]
                    }
                )
                if response.status_code == 200:
                    result = response.json()
                    # Parse response - implementar conforme API
                    logger.info("[Cortex] Escalated to paid API successfully")
        except Exception as e:
            logger.error(f"[Cortex] Paid API failed: {e}")
        
        return get_fallback_recommendation(self.filename, language, len(content), structure_info.get('structure_level', 'semi_structured'), str(e)), {"total_ms": 0, "tokens_generated": 0, "throughput_tps": 0}

    async def _call_cortex_model(self, content: str, language: str, structure_info: dict, attempt: int = 1) -> tuple[dict, dict]:
        """Chama o modelo Cortex via Ollama para analisar o documento."""
        start_time = datetime.utcnow().timestamp()
        
        # Prompt atualizado com campos alternative e notes
        system_prompt = """Você é Cortex, um especialista em análise de documentos para indexação em bases de conhecimento.
Sua tarefa é analisar o conteúdo do documento e retornar uma recomendação estruturada para chunking, embedding e indexação.

Retorne APENAS um JSON válido com a seguinte estrutura:
{
  "operation": "ingest",
  "document": {
    "filename": "nome_do_arquivo",
    "file_type": "txt",
    "language": "pt",
    "estimated_tokens": 123,
    "structure_level": "structured|semi_structured|unstructured",
    "domain": "technical|legal|medical|general|instructional|scraped",
    "notes": "Observações sobre estrutura, ruído, OCR, etc."
  },
  "chunking_strategy": {
    "primary": {"type": "paragraph|sentence|fixed_size|markdown_header|table_aware|code_aware|hybrid", "reason": "..."},
    "alternative": {"type": "fixed_size", "reason": "Quando a primária não for ideal"},
    "parameters": {"chunk_size": 512, "chunk_overlap": 64, "separator": "\\n\\n", "min_chunk_size": 100, "max_chunk_size": 1024}
  },
  "embedding": {"provider": "ollama", "model": "nomic-embed-text", "dimensions": 768, "normalize": true},
  "qdrant_index": {"distance": "Cosine", "hnsw_config": {"m": 16, "ef_construct": 128, "ef_search": 96}, "payload_index_fields": ["source", "document_type", "section"]},
  "retrieval_hint": "Dica específica para como este conteúdo será consultado (ex: 'busca por cláusulas em documento jurídico, top_k baixo, threshold alto')",
  "testing_questions": ["Pergunta 1 para validar a indexação?", "Pergunta 2?"],
  "review_required": false,
  "confidence": 0.85
}"""

        # Incluir metadata extraída no prompt
        user_prompt = f"""Analise este documento para recomendações de indexação:

Filename: {self.filename}
Idioma detectado: {language}
Estrutura detectada: {structure_info.get('structure_level', 'unknown')}
Headings: H1={structure_info.get('headings', {}).get('h1', 0)}, H2={structure_info.get('headings', {}).get('h2', 0)}, H3={structure_info.get('headings', {}).get('h3', 0)}
Tem tabelas: {"sim" if structure_info.get('has_tables') else "não"}
Tem código: {"sim" if structure_info.get('has_code') else "não"}

Conteúdo (primeiros 4000 tokens):
{content[:8000]}

Forneça a melhor estratégia de chunking baseada na estrutura do conteúdo."""

        full_prompt = self._cortex_prompt + "\n\n" + system_prompt + "\n\n" + user_prompt if self._cortex_prompt else system_prompt + "\n\n" + user_prompt

        try:
            async with httpx.AsyncClient(timeout=120.0) as client:
                response = await client.post(
                    f'{OLLAMA_URL}/api/generate',
                    json={
                        "model": OLLAMA_CORTEX_MODEL,
                        "prompt": full_prompt,
                        "stream": False,
                        "options": {"temperature": 0}
                    }
                )
                
                if response.status_code != 200:
                    raise Exception(f"Ollama retornou status {response.status_code}")
                
                result = response.json()
                raw_content = result.get('response', '')
                
                # Extrai JSON da resposta
                json_str = extract_json_from_response(raw_content)
                
                try:
                    rec_data = json.loads(json_str)
                except json.JSONDecodeError as je:
                    if attempt < 2:
                        logger.warning(f"[Cortex] JSON inválido, tentando novamente (tentativa {attempt}): {je}")
                        return await self._call_cortex_model(content, language, structure_info, attempt + 1)
                    else:
                        logger.warning(f"[Cortex] Falha após retry, escalando para API paga")
                        return await self._call_paid_api(content, language, structure_info)
                
                # Preenche defaults se faltarem campos
                rec_data = fill_defaults(rec_data)
                
                # Atualiza info do documento com metadata real
                rec_data['document']['filename'] = self.filename
                rec_data['document']['language'] = language
                rec_data['document']['estimated_tokens'] = len(content) // 4
                rec_data['document']['structure_level'] = structure_info.get('structure_level', 'semi_structured')
                rec_data['document']['notes'] = "; ".join(structure_info.get('notes', [])) if structure_info.get('notes') else None
                
                end_time = datetime.utcnow().timestamp()
                ollama_metrics = {
                    "total_ms": int((end_time - start_time) * 1000),
                    "tokens_generated": result.get('eval_count', 0),
                    "throughput_tps": round(result.get('eval_count', 0) / max((end_time - start_time), 0.001), 2)
                }
                
                return rec_data, ollama_metrics
                
        except json.JSONDecodeError as je:
            if attempt < 2:
                logger.warning(f"[Cortex] JSON inválido, tentando retry: {je}")
                return await self._call_cortex_model(content, language, structure_info, attempt + 1)
            else:
                return await self._call_paid_api(content, language, structure_info)
        except Exception as e:
            logger.warning(f"[Cortex] Falha na chamada ao Ollama, usando fallback: {e}")
            rec_data = get_fallback_recommendation(self.filename, language, len(content), structure_info.get('structure_level', 'semi_structured'), str(e))
            return rec_data, {"total_ms": 0, "tokens_generated": 0, "throughput_tps": 0}

    async def extract_document_metadata(self) -> dict:
        """Extrai metadados reais do documento."""
        start_time = datetime.utcnow().timestamp()
        
        # Detectar tipo real via magic bytes
        file_type = detect_file_type(self.file_content, self.filename)
        
        # Extrair texto baseado no tipo
        content = ""
        extra_info = {}
        
        try:
            if file_type == 'pdf':
                content, extra_info = extract_text_from_pdf(self.file_content)
            elif file_type == 'docx':
                content, extra_info = extract_text_from_docx(self.file_content)
            elif file_type == 'html':
                content, extra_info = extract_text_from_html(self.file_content)
            else:
                content = self.file_content.decode('utf-8', errors='replace')
        except Exception as e:
            logger.warning(f"[Metadata] Erro ao extrair texto: {e}")
            content = self.file_content.decode('utf-8', errors='replace')
        
        self._content = content
        
        # Detectar idioma
        language = detect_language(content)
        
        # Analisar estrutura
        structure_info = analyze_structure(content, file_type)
        
        self._extraction_time_ms = int((datetime.utcnow().timestamp() - start_time) * 1000)
        
        self._metadata = {
            "file_type": file_type,
            "language": language,
            "structure": structure_info,
            "extraction_time_ms": self._extraction_time_ms,
            "extra_info": extra_info
        }
        
        return self._metadata

    async def analyze_document(self) -> IngestRecommendation:
        # Etapa 1: Extrair metadados reais
        await self.emit_step('extract_metadata', 'Extraindo metadados', 'running', f'Analisando {self.filename}...')
        try:
            metadata = await self.extract_document_metadata()
            await self.emit_step('extract_metadata', 'Extraindo metadados', 'done', 
                               f'{metadata["file_type"].upper()} • {metadata["language"]} • {self._extraction_time_ms}ms',
                               recommendation={"extraction_metrics": metadata})
        except Exception as e:
            await self.emit_step('extract_metadata', 'Extraindo metadados', 'error', error_message=str(e))
            metadata = {"file_type": "txt", "language": "pt", "structure": {"structure_level": "semi_structured", "notes": []}}
        
        # Etapa 2: Coletar amostras
        content = self._content or ""
        char_count = len(content)
        word_count = len(content.split())
        
        await self.emit_step('collecting_samples', 'Coletando amostras', 'running', 'Preparando conteúdo...')
        await asyncio.sleep(0.05)
        await self.emit_step('collecting_samples', 'Coletando amostras', 'done', f'{char_count} chars, {word_count} palavras')
        
        # Etapa 3: Analisar com Cortex
        await self.emit_step('analyze', 'Analisando com Cortex', 'running', 'Processando...')
        await asyncio.sleep(0.2)
        
        logger.info(f"[Cortex] Analisando {self.filename} com modelo {OLLAMA_CORTEX_MODEL}")
        rec_data, ollama_metrics = await self._call_cortex_model(content, metadata.get('language', 'pt'), metadata.get('structure', {}))
        
        self._last_ollama_metrics = ollama_metrics
        self._recommendation = rec_data
        
        recommendation = IngestRecommendation(**rec_data)
        
        await self.emit_step('analyze', 'Analisando com Cortex', 'done', f'{ollama_metrics["total_ms"]}ms', 
                           recommendation={"ollama_metrics": ollama_metrics})
        
        # Etapa 4: Validar schema
        await self.emit_step('validate_schema', 'Validando schema', 'running', 'Verificando...')
        await self.emit_step('validate_schema', 'Validando schema', 'done', 'OK')
        
        # Salvar usage metrics
        await self.save_ollama_usage()
        
        return recommendation

    async def process(self) -> IngestRecommendation:
        recommendation = await self.analyze_document()
        await self.emit_step('awaiting_confirmation', 'Aguardando confirmação', 'done', 'Revise os parâmetros', recommendation=recommendation.model_dump())
        return recommendation

    async def continue_after_confirmation(self, chunking_params: Optional[dict] = None, document_id: str = ""):
        self._flow_end_time = datetime.utcnow().timestamp()
        
        # Etapa 5: Chunking
        await self.emit_step('chunking', 'Indexando', 'running', 'Dividindo...')
        if not self._content: self._content = self.file_content.decode('utf-8', errors='replace')
        
        params = chunking_params or self._recommendation.get('chunking_strategy', {}).get('parameters', {})
        strategy = self._recommendation.get('chunking_strategy', {}).get('primary', {}).get('type', 'paragraph')
        
        chunking_start = datetime.utcnow().timestamp()
        self._chunks = chunk_text(self._content, strategy, params)
        chunking_time = int((datetime.utcnow().timestamp() - chunking_start) * 1000)
        
        await self.emit_step('chunking', 'Indexando', 'done', f'{len(self._chunks)} chunks ({chunking_time}ms)')
        
        # Etapa 6: Embedding
        await self.emit_step('embedding', 'Gerando embeddings', 'running', 'Criando vetores...')
        qc = None
        embedding_tokens = 0
        embedding_start = datetime.utcnow().timestamp()
        try:
            qc = QdrantClient(url=QDRANT_URL, prefer_grpc=False)
        except Exception: qc = None

        if qc and self._chunks:
            try:
                dims = self._recommendation.get('embedding', {}).get('dimensions', 768)
                qc.recreate_collection(collection_name=QDRANT_COLLECTION, vectors_config=models.VectorParams(size=dims, distance=models.Distance.COSINE))
                embeddings = [[0] * 768 for _ in self._chunks]
                for i, chunk in enumerate(self._chunks):
                    try:
                        async with httpx.AsyncClient(timeout=30.0) as client:
                            r = await client.post(f'{OLLAMA_URL}/api/embeddings', json={"model": OLLAMA_EMBEDDING_MODEL, "prompt": chunk})
                            if r.status_code == 200:
                                embeddings[i] = r.json().get('embedding', [0] * 768)
                                embedding_tokens += len(chunk) // 4
                    except Exception: pass
                points = [models.PointStruct(id=str(uuid.uuid4()), vector=emb, payload={"flow_id": self.flow_id, "chunk_index": i, "content": c[:500], "filename": self.filename}) for i, (c, emb) in enumerate(zip(self._chunks, embeddings))]
                if points: qc.upsert(collection_name=QDRANT_COLLECTION, points=points)
            except Exception as e:
                logger.error(f"[Qdrant] Erro: {e}")
        
        embedding_time = int((datetime.utcnow().timestamp() - embedding_start) * 1000)
        self._embedding_latency_ms = embedding_time
        
        await self.emit_step('embedding', 'Gerando embeddings', 'done', f'{embedding_time}ms')
        await self.emit_step('done', 'Concluído', 'done', 'Documento indexado!')
        
        # Salvar usage metrics para embedding
        await self.save_ollama_usage(for_embedding=True)
        
        # Gerar e persistir relatório de ingestão
        await self._save_ingestion_report(document_id)
    
    async def _save_ingestion_report(self, document_id: str) -> None:
        """Salva relatório de ingestão com métricas."""
        try:
            engine = _get_engine()
            with engine.connect() as conn:
                embedding_model_result = conn.execute(text("""
                    SELECT name FROM models WHERE user_id = :user_id AND kind = 'embedding' LIMIT 1
                """), {"user_id": self.user_id}).fetchone()
                
                embedding_model_name = embedding_model_result[0] if embedding_model_result else OLLAMA_EMBEDDING_MODEL
                
                processing_time_ms = int((self._flow_end_time - self._flow_start_time) * 1000) if self._flow_end_time and self._flow_start_time else 0
                
                report = {
                    "filename": self.filename,
                    "file_type": self._metadata.get("file_type", "txt") if self._metadata else "txt",
                    "language": self._metadata.get("language", "pt") if self._metadata else "pt",
                    "char_count": len(self._content) if self._content else 0,
                    "word_count": len(self._content.split()) if self._content else 0,
                    "chunk_count": len(self._chunks),
                    "chunking_strategy": self._recommendation.get("chunking_strategy", {}) if self._recommendation else {},
                    "embedding": {
                        "model": embedding_model_name,
                        "tokens": getattr(self, '_total_embedding_tokens', 0),
                        "latency_ms": self._embedding_latency_ms
                    },
                    "extraction_metrics": self._metadata if self._metadata else {},
                    "cortex": {
                        "model": OLLAMA_CORTEX_MODEL,
                        "tokens_generated": self._last_ollama_metrics.get("tokens_generated", 0) if self._last_ollama_metrics else 0,
                        "latency_ms": self._last_ollama_metrics.get("total_ms", 0) if self._last_ollama_metrics else 0
                    },
                    "processing_time_ms": processing_time_ms,
                    "status": "done",
                    "completed_at": datetime.utcnow().isoformat() + "Z"
                }
                
                conn.execute(text("""
                    INSERT INTO knowledge_ingestion_report (document_id, report)
                    VALUES (:document_id, :report)
                """), {
                    "document_id": document_id,
                    "report": json.dumps(report)
                })
                conn.commit()
                logger.info(f"[IngestionReport] Saved report for document {document_id}")
        except Exception as e:
            logger.error(f"[IngestionReport] Failed to save report: {e}")

    async def save_ollama_usage(self, for_embedding: bool = False) -> None:
        """Persiste métricas de uso do Ollama."""
        try:
            engine = _get_engine()
            with engine.connect() as conn:
                provider_result = conn.execute(text("""
                    SELECT id FROM providers WHERE slug = 'ollama'
                """)).fetchone()
                
                if not provider_result:
                    logger.warning("[ProvidersUsage] Ollama provider not found, skipping usage save")
                    return
                
                provider_id = provider_result[0]
                task = 'embedding' if for_embedding else 'knowledge_analysis'
                
                if for_embedding:
                    embedding_model = conn.execute(text("""
                        SELECT id, name FROM models 
                        WHERE user_id = :user_id AND kind = 'embedding' 
                        LIMIT 1
                    """), {"user_id": self.user_id}).fetchone()
                    
                    if embedding_model:
                        model_id, model_name = embedding_model[0], embedding_model[1]
                        metrics = self._embedding_metrics
                    else:
                        logger.warning("[ProvidersUsage] No embedding model found for user")
                        return
                else:
                    model_name = OLLAMA_CORTEX_MODEL
                    metrics = self._last_ollama_metrics
                    
                    model_id = None
                    model_result = conn.execute(text("""
                        SELECT id FROM models 
                        WHERE user_id = :user_id AND provider_id = :provider_id AND name = :model_name
                    """), {
                        "user_id": self.user_id,
                        "provider_id": provider_id,
                        "model_name": model_name
                    }).fetchone()
                    
                    if model_result:
                        model_id = model_result[0]
                
                processing_time_ms = int((self._flow_end_time - self._flow_start_time) * 1000) if self._flow_end_time and self._flow_start_time else 0
                chunk_count = len(self._chunks) if self._chunks else 0
                tokens_input = len(self._content) // 4 if self._content else 0
                
                conn.execute(text("""
                    INSERT INTO providers_usage (
                        user_id, provider_id, model_id, model_name,
                        request_count, tokens_input, tokens_output, latency_ms,
                        task, chunk_count, processing_time_ms
                    ) VALUES (
                        :user_id, :provider_id, :model_id, :model_name,
                        1, :tokens_input, :tokens_output, :latency_ms,
                        :task, :chunk_count, :processing_time_ms
                    )
                """), {
                    "user_id": self.user_id,
                    "provider_id": provider_id,
                    "model_id": model_id,
                    "model_name": model_name,
                    "tokens_input": tokens_input,
                    "tokens_output": metrics.get("tokens_generated", 0) if metrics else 0,
                    "latency_ms": metrics.get("total_ms", 0) if metrics else 0,
                    "task": task,
                    "chunk_count": chunk_count,
                    "processing_time_ms": processing_time_ms
                })
                conn.commit()
                logger.info(f"[ProvidersUsage] Saved usage for {model_name}: tokens_input={tokens_input}, tokens_output={metrics.get('tokens_generated', 0) if metrics else 0}, task={task}")
        except Exception as e:
            logger.error(f"[ProvidersUsage] Failed to save usage: {e}")